from __future__ import annotations

import json
import re
import subprocess
from collections import OrderedDict
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "TAZA041_Technical_Operations_API_Guide_AR.docx"
LOGO = ROOT / "presentation_assets" / "taza041-logo.jpg"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "172033"
MUTED = "5D6778"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"
GREEN = "237A57"
AMBER = "7A5A00"
RED = "9B1C1C"
FONT = "Arial"  # named preset override: clearer Arabic shaping than Calibri.
MONO = "Consolas"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
    cant_split.set(qn("w:val"), "1")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    total = sum(widths_dxa)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        set_row_cant_split(row)
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color="C8D0DA", size="6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_rtl(paragraph, rtl=True) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if rtl:
        if bidi is None:
            bidi = OxmlElement("w:bidi")
            p_pr.append(bidi)
        bidi.set(qn("w:val"), "1")
    elif bidi is not None:
        p_pr.remove(bidi)


def set_run_font(run, name=FONT, size=None, bold=None, color=None, italic=None) -> None:
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{key}"), name)
    if size is not None:
        run.font.size = Pt(size)
        sz_cs = r_pr.find(qn("w:szCs"))
        if sz_cs is None:
            sz_cs = OxmlElement("w:szCs")
            r_pr.append(sz_cs)
        sz_cs.set(qn("w:val"), str(int(size * 2)))
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_page_number(paragraph) -> None:
    set_rtl(paragraph, False)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl(normal.element.get_or_add_pPr(), True) if False else None

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[style_name]
        style.font.name = FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        p_pr = style.element.get_or_add_pPr()
        bidi = p_pr.find(qn("w:bidi"))
        if bidi is None:
            bidi = OxmlElement("w:bidi")
            p_pr.append(bidi)

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = FONT
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.right_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        p_pr = style.element.get_or_add_pPr()
        bidi = p_pr.find(qn("w:bidi"))
        if bidi is None:
            bidi = OxmlElement("w:bidi")
            p_pr.append(bidi)

    if "Code Block" not in styles:
        code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = styles["Code Block"]
    code.font.name = MONO
    code.font.size = Pt(8.5)
    code.font.color.rgb = RGBColor.from_string(INK)
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.right_indent = Inches(0.18)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(6)
    code.paragraph_format.line_spacing = 1.0


def add_body(doc, text, *, bold_label=None, keep=False, color=None):
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.keep_together = keep
    if bold_label and text.startswith(bold_label):
        r1 = p.add_run(bold_label)
        set_run_font(r1, bold=True, color=color or INK)
        r2 = p.add_run(text[len(bold_label):])
        set_run_font(r2, color=color or INK)
    else:
        r = p.add_run(text)
        set_run_font(r, color=color or INK)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    set_run_font(run)
    return p


def new_numbering(doc) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "right")
    p_pr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:right"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    p_pr.append(bidi)
    lvl.extend([start, num_fmt, lvl_text, lvl_jc, p_pr])
    abstract.append(lvl)
    first_num_index = next(
        (index for index, child in enumerate(numbering) if child.tag == qn("w:num")),
        len(numbering),
    )
    numbering.insert(first_num_index, abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_numbered(doc, text, num_id=None):
    p = doc.add_paragraph(style="List Number")
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if num_id is not None:
        p_pr = p._p.get_or_add_pPr()
        num_pr = p_pr.find(qn("w:numPr"))
        if num_pr is None:
            num_pr = OxmlElement("w:numPr")
            p_pr.append(num_pr)
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id_node = OxmlElement("w:numId")
        num_id_node.set(qn("w:val"), str(num_id))
        num_pr.extend([ilvl, num_id_node])
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(text)
    set_run_font(r, size={1: 16, 2: 13, 3: 12}[level], bold=True,
                 color={1: BLUE, 2: BLUE, 3: DARK_BLUE}[level])
    return p


def add_callout(doc, title, text, fill=LIGHT_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, [9360])
    set_table_borders(table, color=accent, size="8")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title + " — ")
    set_run_font(r, bold=True, color=accent)
    r = p.add_run(text)
    set_run_font(r, color=INK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    return table


def add_label_detail_table(doc, rows, widths=(2700, 6660), header=None):
    total_rows = len(rows) + (1 if header else 0)
    table = doc.add_table(rows=total_rows, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, list(widths))
    set_table_borders(table)
    offset = 0
    if header:
        set_repeat_table_header(table.rows[0])
        for i, value in enumerate(header):
            set_cell_shading(table.cell(0, i), LIGHT_BLUE)
            p = table.cell(0, i).paragraphs[0]
            set_rtl(p)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            set_run_font(p.add_run(value), bold=True, color=DARK_BLUE, size=10)
        offset = 1
    elif table.rows:
        # Accessibility: the first key/value row acts as the table's identifying row.
        set_repeat_table_header(table.rows[0])
    for r_idx, (label, detail) in enumerate(rows, start=offset):
        set_cell_shading(table.cell(r_idx, 0), LIGHT_GRAY)
        for c_idx, value in enumerate((label, detail)):
            p = table.cell(r_idx, c_idx).paragraphs[0]
            set_rtl(p)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            set_run_font(p.add_run(value), bold=(c_idx == 0), size=9.5,
                         color=DARK_BLUE if c_idx == 0 else INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_flow(doc, operation):
    add_heading(doc, operation["name"], 3)
    add_body(doc, "نقطة البدء: " + operation["trigger"], bold_label="نقطة البدء:")
    num_id = new_numbering(doc)
    add_numbered(doc, operation["client"], num_id)
    add_numbered(doc, operation["request"], num_id)
    add_numbered(doc, operation["server"], num_id)
    add_numbered(doc, operation["result"], num_id)
    if operation.get("errors"):
        add_body(doc, "مسارات الفشل: " + operation["errors"], bold_label="مسارات الفشل:", color=RED)


def add_code(doc, text):
    p = doc.add_paragraph(style="Code Block")
    set_rtl(p, False)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F6F8FA")
    p_pr.append(shd)
    r = p.add_run(text)
    set_run_font(r, name=MONO, size=8.5)
    return p


def load_routes():
    proc = subprocess.run(
        ["php", "artisan", "route:list", "--path=api", "--json"],
        cwd=ROOT,
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
        check=True,
    )
    return json.loads(proc.stdout)


ACTION_PURPOSES = {
    "login": "يتحقق من بيانات الدخول وينشئ توكن Sanctum جديدًا.",
    "logout": "يحذف توكن الجلسة الحالي وينهي المصادقة.",
    "me": "يعيد ملف الموظف الحالي وإحصاءات مرتبطة بدوره.",
    "register": "ينشئ حساب زبون وتوكنًا وحساب ولاء ابتدائيًا.",
    "forgotPassword": "يطلب رابط استعادة مع استجابة عامة لا تكشف وجود الحساب.",
    "resetPassword": "يتحقق من الرمز ويغيّر كلمة المرور ويلغي الجلسات القديمة.",
    "profile": "يعيد ملف الزبون وبياناته المرتبطة.",
    "updateProfile": "يتحقق من المدخلات وكلمة المرور عند اللزوم ثم يحفظ التعديلات.",
    "index": "يعيد قائمة مفلترة ومقسمة إلى صفحات بحسب سياق المسار.",
    "show": "يعيد تفاصيل السجل المطلوب بعد التحقق من الصلاحية والوجود.",
    "store": "يتحقق من البيانات وينشئ سجلًا جديدًا ويرجع تفاصيله.",
    "update": "يتحقق من البيانات ويحدّث السجل ثم يرجع النسخة الجديدة.",
    "destroy": "يتحقق من قابلية الحذف ثم يحذف السجل أو يرفض العملية.",
    "active": "يعيد السجلات النشطة حاليًا.",
    "expired": "يعيد السجلات المنتهية.",
    "upcoming": "يعيد السجلات القادمة.",
    "stats": "يجمع مؤشرات وإحصاءات الشاشة.",
    "publicIndex": "يعيد البيانات المتاحة للعرض العام فقط.",
    "publicShow": "يعيد تفاصيل عنصر متاح للعامة.",
    "liveData": "يعيد لقطة موحدة لبيانات الموقع الحية لتقليل الطلبات.",
    "publicInfo": "يعيد معلومات المطعم العامة.",
    "publicImages": "يعيد صور المطعم الفعالة والمرتبة.",
    "pricingInfo": "يعيد إعدادات أسعار التوصيل والحجز والولاء.",
    "deliveryQuote": "يحسب طريق التوصيل والمسافة والتكلفة وحد النطاق.",
    "chat": "يعالج رسالة المساعد ويحفظ السياق ويعيد الرد والاقتراحات.",
    "customerHistory": "يعيد آخر محادثات الزبون مجمعة حسب اليوم.",
    "conversations": "يعرض محادثات المساعد للإدارة مع الفلاتر.",
    "conversationStats": "يعيد إحصاءات المحادثات والنوايا.",
    "generateDailyReport": "يضع مهمة تقرير الذكاء الاصطناعي في الطابور.",
    "customerStore": "ينشئ طلبًا مع عناصره ومرفق التوصيل أو الحجز داخل معاملة.",
    "customerOrders": "يعيد طلبات الزبون وحالاتها الموحدة.",
    "customerShow": "يعيد طلبًا يخص الزبون الحالي فقط.",
    "customerCancel": "يلغي الطلب المسموح ويسوي الدفع والمخزون والنقاط مرة واحدة.",
    "customerPay": "ينشئ سجل الدفع ويعالج الطريقة ويحدّث الولاء والإشعارات.",
    "changeStatus": "يتحقق من انتقال الحالة والدور ثم يحدّث السجل وآثاره.",
    "notifyCustomer": "ينشئ إشعارًا يدويًا للزبون المرتبط بالسجل.",
    "archive": "ينقل السجل التشغيلي المغلق إلى الأرشيف.",
    "restore": "يعيد السجل المؤرشف إلى العرض.",
    "restoreArchive": "يعيد سجل الطلب من الأرشيف.",
    "pending": "يعيد العناصر المعلقة التي تتطلب إجراءً.",
    "normalOrders": "يعيد الطلبات العادية فقط.",
    "normalStats": "يعيد مؤشرات الطلبات العادية.",
    "adminIndex": "يعيد عرضًا إداريًا شاملاً ومفلترًا.",
    "adminShow": "يعيد التفاصيل الموسعة للمدير العام.",
    "today": "يعيد حجوزات اليوم.",
    "tables": "يعيد كتالوج الطاولات وتوفرها وتسعيرها.",
    "tableAvailability": "يفحص تعارضات الطاولة في الوقت المحدد.",
    "assigned": "يعيد عمليات التوصيل المسندة للسائق الحالي.",
    "getDrivers": "يعيد السائقين النشطين وحالة انشغالهم.",
    "assignDriver": "يتحقق من جاهزية الطلب والسائق ثم يسند التوصيل ويرسل إشعارات.",
    "driverRatings": "يعيد تقييمات السائق وتوزيعها.",
    "driverStats": "يعيد مؤشرات أداء السائق.",
    "settings": "يعيد إعدادات الوحدة الحالية.",
    "updateSettings": "يتحقق من إعدادات الوحدة ويحفظها.",
    "lowStock": "يعيد المنتجات منخفضة المخزون.",
    "outOfStock": "يعيد المنتجات النافدة.",
    "updatePrice": "يعدّل سعر المنتج بعد التحقق.",
    "updateLoyaltyPrice": "يعدّل سعر الاستبدال بنقاط الولاء.",
    "updateStock": "يعدّل كمية المخزون ويطلق التنبيهات اللازمة.",
    "toggle": "يبدّل حالة التفعيل/التوفر.",
    "reportUnavailable": "يسجل بلاغ الزبون ويرسل تنبيهًا لمدير المخزون.",
    "addProduct": "يربط منتجًا بالعرض مع الكمية.",
    "removeProduct": "يفصل منتجًا عن العرض.",
    "broadcast": "يضع إرسال إشعار جماعي في الطابور.",
    "accountsIndex": "يعرض حسابات التحصيل.",
    "accountsSummary": "يلخص الأرصدة والسعات والحساب الأساسي.",
    "accountShow": "يعيد تفاصيل حساب تحصيل.",
    "accountStore": "ينشئ حساب تحصيل جديدًا.",
    "accountUpdate": "يعدّل بيانات حساب التحصيل.",
    "updateBalance": "يعدّل الرصيد إداريًا مع تسجيل الأثر.",
    "makePrimary": "يجعل الحساب المحدد أساسيًا لنوعه.",
    "withdraw": "ينفذ سحبًا داخليًا بعد فحص الرصيد.",
    "accountDestroy": "يحذف حسابًا مسموحًا حذفه.",
    "paymentsIndex": "يعرض حركات الدفع المفلترة.",
    "paymentShow": "يعيد سجل دفع مع الطلب والزبون.",
    "refund": "يعكس الدفع ويحدّث حالة السجل.",
    "generateReport": "يضع توليد التقرير المالي في الطابور.",
    "customerAccount": "يعيد رصيد ومستوى وحركات ولاء الزبون.",
    "transactions": "يعيد حركات الولاء للزبون المحدد.",
    "adjust": "يضيف أو يخصم نقاطًا إداريًا ويسجل الحركة.",
    "employeeIndex": "يعيد إشعارات الموظف الحالي.",
    "employeeUnreadCount": "يعيد عدد إشعارات الموظف غير المقروءة.",
    "employeeMarkAllRead": "يعلّم كل إشعارات الموظف كمقروءة.",
    "customerIndex": "يعيد السجلات الخاصة بالزبون الحالي.",
    "customerMarkAllRead": "يعلّم كل إشعارات الزبون كمقروءة.",
    "markRead": "يعلّم الإشعار المحدد كمقروء إذا كان للمستخدم.",
    "markReviewed": "يعلّم التقرير كمراجع.",
    "adminSend": "يرسل تعليمات المدير المرتبطة بالتقرير.",
    "aiReportsIndex": "يعرض تقارير الذكاء الاصطناعي لمدير التواصل.",
    "forwardToGM": "يمرر تقرير الذكاء الاصطناعي للمدير العام.",
    "commShow": "يعيد معلومات المطعم القابلة للتحرير لمدير التواصل.",
    "updateContactInfo": "يحدّث معلومات الموقع والتواصل والمحتوى العام.",
    "imagesIndex": "يعيد معرض صور المطعم.",
    "imageUpdate": "يحدّث وصف الصورة وتصنيفها.",
    "imageReorder": "يغير ترتيب الصورة.",
    "imageToggle": "يبدّل ظهور الصورة للعامة.",
    "imageDestroy": "يحذف سجل الصورة وملفها.",
    "review": "يضيف ملاحظة المراجعة ويغيّر حالة الاقتراح.",
    "markImplemented": "يعلّم اقتراح الوجبة كمطبق ويبلغ الزبون.",
    "reject": "يرفض الاقتراح مع السبب ويبلغ الزبون.",
    "rateEmployee": "يحفظ تقييم الموظف وفق الصلاحية.",
    "employeeReviews": "يعيد تقييمات الموظفين أو موظف محدد.",
    "driverReviews": "يعيد تقييمات السائقين.",
    "driverSummary": "يعيد ملخص تقييم السائق.",
    "customerProductReviews": "يعرض تقييمات الوجبات الصادرة عن الزبائن.",
    "customerRateDriver": "يحفظ تقييم السائق بعد التسليم.",
    "customerRateProduct": "يحفظ تقييم وجبة بعد اكتمال الطلب.",
    "uploadProductImage": "يتحقق من الصورة ويخزنها ويربطها بالمنتج.",
    "uploadOfferImage": "يتحقق من الصورة ويخزنها ويربطها بالعرض.",
    "uploadRestaurantImage": "يرفع صورة للمعرض مع بياناتها.",
    "uploadLogo": "يرفع شعار المطعم ويحدّث الإعدادات.",
    "uploadEmployeeAvatar": "يرفع صورة الموظف بعد التحقق من الصلاحية.",
    "deleteEmployeeAvatar": "يحذف صورة الموظف ويفرغ الحقل.",
    "uploadCustomerAvatar": "يرفع صورة الزبون مع تأكيد كلمة المرور.",
    "uploadImage": "يرفع صورة عامة إلى المسار المسموح.",
    "deleteFile": "يحذف ملفًا من التخزين ضمن القيود.",
    "sync": "يستبدل مجموعة العناوين المحفوظة بعد التحقق.",
}


def route_group(uri: str) -> tuple[str, str]:
    path = uri.removeprefix("api/")
    if path == "health": return ("الصحة والجاهزية", "عام")
    if path.startswith("public/"): return ("الواجهات العامة", "عام")
    if path.startswith("customer/auth/"): return ("مصادقة الزبون", "عام/زبون")
    if path.startswith("customer/"): return ("عمليات الزبون المحمية", "زبون")
    if path.startswith("auth/employee/"): return ("مصادقة الموظف", "موظف")
    if path.startswith("employee/"): return ("العمليات المشتركة للموظفين", "موظف")
    if path.startswith("admin/"): return ("المدير العام", "general_manager")
    if path.startswith("orders/") or path == "orders": return ("إدارة الطلبات والحجوزات", "order_manager")
    if path.startswith("delivery/") or path == "delivery": return ("التوصيل والسائقون", "delivery_manager / driver")
    if path.startswith("products/") or path == "products": return ("المنتجات والمخزون", "inventory_manager")
    if path.startswith("offers/") or path == "offers": return ("العروض", "inventory_manager")
    if path.startswith("finance/"): return ("المالية والمدفوعات", "finance_manager")
    if path.startswith("loyalty/") or path == "loyalty": return ("الولاء", "general_manager")
    if path.startswith("communication/"): return ("التواصل والمحتوى", "communication_manager")
    if path.startswith("ai/"): return ("إدارة الذكاء الاصطناعي", "communication_manager / general_manager")
    if path.startswith("reviews/"): return ("التقييمات", "موظف مصرح")
    if path.startswith("upload/"): return ("الرفع العام للموظفين", "موظف مصرح")
    return ("مسارات أخرى", "بحسب المتحكم")


def route_purpose(route) -> str:
    action = route.get("action", "")
    if action == "Closure":
        return "يفحص اتصال التطبيق وقاعدة البيانات ويعيد حالة الجاهزية."
    method_name = action.split("@")[-1]
    return ACTION_PURPOSES.get(method_name, "ينفذ عملية المسار بعد التحقق من الصلاحية والمدخلات.")


def add_route_catalog(doc, routes):
    grouped = OrderedDict()
    for route in routes:
        group = route_group(route["uri"])
        grouped.setdefault(group, []).append(route)

    for (group_name, access), items in grouped.items():
        add_heading(doc, group_name, 2)
        add_body(doc, f"الوصول المعتاد: {access}. عدد المسارات في هذه المجموعة: {len(items)}.", bold_label="الوصول المعتاد:")
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        set_table_geometry(table, [780, 4230, 4350])
        set_table_borders(table)
        headers = ("الطريقة", "المسار", "المتحكم والغرض")
        set_repeat_table_header(table.rows[0])
        for idx, label in enumerate(headers):
            cell = table.cell(0, idx)
            set_cell_shading(cell, LIGHT_BLUE)
            p = cell.paragraphs[0]
            set_rtl(p)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.RIGHT
            set_run_font(p.add_run(label), bold=True, color=DARK_BLUE, size=9)
        for route in items:
            row = table.add_row()
            set_row_cant_split(row)
            method = route["method"].replace("|HEAD", "")
            uri = "/" + route["uri"]
            action = route["action"].replace("App\\Http\\Controllers\\API\\", "")
            values = (method, uri, f"{action} — {route_purpose(route)}")
            for idx, value in enumerate(values):
                p = row.cells[idx].paragraphs[0]
                if idx == 1:
                    set_rtl(p, False)
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    set_run_font(p.add_run(value), name=MONO, size=7.4, color=INK)
                else:
                    set_rtl(p)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.RIGHT
                    set_run_font(p.add_run(value), name=MONO if idx == 0 else FONT, size=7.6,
                                 bold=(idx == 0), color=DARK_BLUE if idx == 0 else INK)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)


def main():
    routes = load_routes()
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)

    header = section.header
    hp = header.paragraphs[0]
    set_rtl(hp)
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(hp.add_run("TAZA 041  |  دليل العمليات وواجهات API"), size=8.5, color=MUTED, bold=True)
    add_page_number(section.footer.paragraphs[0])

    # Cover: editorial_cover pattern, compact and technical.
    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(LOGO), width=Inches(1.05))
        inline = run._r.xpath(".//wp:inline")[0]
        doc_pr = inline.find(qn("wp:docPr"))
        if doc_pr is not None:
            doc_pr.set("descr", "شعار مطعم TAZA 041")
            doc_pr.set("title", "TAZA 041")
        doc.add_paragraph().paragraph_format.space_after = Pt(20)
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    p.paragraph_format.space_after = Pt(8)
    set_run_font(p.add_run("الدليل التقني الشامل للعمليات"), size=28, bold=True, color=DARK_BLUE)
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    set_run_font(p.add_run("من ضغط الزر في الواجهة إلى تنفيذ Laravel وإرجاع النتيجة"), size=15, color=BLUE)
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(30)
    set_run_font(p.add_run("موقع الزبون • لوحة الموظفين • 188 مسار API • المهام الخلفية والخدمات الخارجية"), size=10.5, color=MUTED, bold=True)
    add_label_detail_table(doc, [
        ("المشروع", "TAZA 041 — نظام إدارة مطعم"),
        ("الغرض", "مرجع للمطورين والمختبرين وفرق التسليم والصيانة"),
        ("نطاق المصدر", "المستودع الجذري الحالي؛ لا تعتمد نسخة TAZA041_Delivery_2026-08-20 المكررة كمصدر ثانٍ"),
        ("تاريخ اللقطة", date.today().isoformat()),
        ("إعداد بصري", "compact_reference_guide مع تجاوز موثق: خط Arial لتحسين تشكيل العربية"),
    ])
    add_callout(doc, "طريقة القراءة", "لكل عملية مهمة يوضح الدليل نقطة الضغط، منطق الواجهة، الطلب، المتحكم والخدمة، أثر قاعدة البيانات، شكل النجاح، ومسارات الخطأ. وفي النهاية ملحق كامل مولّد من قائمة المسارات الفعلية.", fill="EEF6F2", accent=GREEN)
    doc.add_page_break()

    add_heading(doc, "فهرس المحتويات", 1)
    for item in (
        "1. حدود التوثيق وخريطة النظام",
        "2. دورة الطلب المشتركة والاستجابات والمصادقة",
        "3. عمليات موقع الزبون طرفًا إلى طرف",
        "4. العمليات المشتركة في لوحة الموظفين",
        "5. عمليات المدير العام",
        "6. عمليات مدير الطلبات والحجوزات",
        "7. عمليات مدير التوصيل والسائق",
        "8. عمليات مدير المخزون والعروض",
        "9. عمليات المدير المالي",
        "10. عمليات مدير التواصل والذكاء الاصطناعي",
        "11. المهام الخلفية والخدمات الخارجية والنسخ الاحتياطي",
        "12. خرائط الحالات والآثار المتقاطعة",
        "13. ملحق API الكامل (188 مسارًا)",
        "14. دليل الملفات المرجعية",
    ):
        add_bullet(doc, item)

    add_heading(doc, "1. حدود التوثيق وخريطة النظام", 1)
    add_body(doc, "يتكون النظام من Laravel API في الخادم، موقع زبون ثابت داخل public/frontend، ولوحات موظفين ثابتة داخل public/dashboard. الواجهات لا تتصل بقاعدة البيانات مباشرة؛ كل البيانات الحية تمر عبر /api، بينما تحفظ الواجهة بيانات جلسة محدودة مثل التوكن والمستخدم والسلة واللغة في Local Storage أو Session Storage.")
    add_label_detail_table(doc, [
        ("موقع الزبون", "public/frontend/*.html وpublic/frontend/assets/js؛ التصفح، السلة، الطلب، الحجز، التوصيل، الدفع، الحساب والاقتراحات."),
        ("لوحة الموظفين", "public/dashboard/*.html وpublic/dashboard/assets/js؛ صفحة مستقلة لكل دور مع عميل HTTP موحد."),
        ("طبقة API", "routes/api.php وapp/Http/Controllers/API؛ مصادقة Sanctum واستجابات JSON موحدة."),
        ("منطق المجال", "app/Models وapp/Services؛ انتقالات الحالات، التسعير، المخزون، الولاء، الإلغاء، والمسارات."),
        ("المهام الخلفية", "app/Jobs وقنوات الطابور mail, notifications, reports, default."),
        ("التخزين", "قاعدة البيانات، storage/app/public للصور، ومجلد النسخ الاحتياطية حسب config/backup.php."),
    ], header=("الطبقة", "المسؤولية"))
    add_callout(doc, "قاعدة المصدر", "كل عدد مسار وكل اسم endpoint في هذا المستند مأخوذ من php artisan route:list --path=api --json وقت إنشاء الملف. الشرح الوظيفي دُقق مقابل المتحكمات والخدمات وملفات JavaScript الجذرية.")

    add_heading(doc, "2. دورة الطلب المشتركة والاستجابات والمصادقة", 1)
    add_heading(doc, "2.1 من الزر إلى JSON", 2)
    request_cycle_num = new_numbering(doc)
    for step in (
        "يربط ملف الصفحة حدث click أو submit بدالة؛ تمنع الدالة الإرسال الافتراضي وتتحقق من المدخلات محليًا وتغيّر الزر إلى حالة مشغول عند الحاجة.",
        "يبني موقع الزبون الطلب عبر apiFetch، وتبني اللوحة الطلب عبر TAZA.Http. يضاف Accept: application/json، وContent-Type: application/json للبيانات العادية، وBearer Token للمسارات المحمية.",
        "يطابق Laravel المسار في routes/api.php ويطبق middleware مثل auth:sanctum وTrackCustomerIp وقيود throttle.",
        "يتحقق المتحكم من نوع المستخدم والدور ثم ينفذ validation. الأخطاء تعاد عادة بكود 422، وعدم الصلاحية 403، وعدم الوجود 404.",
        "تنفذ المتحكمات عمليات قاعدة البيانات، وتستخدم DB transaction في العمليات المركبة مثل إنشاء الطلب والدفع والإلغاء وتغيير بعض الحالات.",
        "تعيد BaseController غلافًا موحدًا: success وmessage وdata عند النجاح؛ أو success=false وmessage وerrors عند الفشل.",
        "يفك عميل الواجهة JSON. عند النجاح يحدث الحالة المحلية والعناصر المرئية ويعرض Toast أو يعيد التوجيه؛ وعند الفشل يعيد تفعيل الزر ويعرض رسالة الخادم أو رسالة آمنة للمستخدم.",
    ):
        add_numbered(doc, step, request_cycle_num)
    add_code(doc, '{\n  "success": true,\n  "message": "تمت العملية بنجاح",\n  "data": { "...": "..." }\n}\n\n{\n  "success": false,\n  "message": "أول رسالة تحقق مفهومة",\n  "errors": { "field": ["تفاصيل الخطأ"] }\n}')

    add_heading(doc, "2.2 فروق عميلَي HTTP", 2)
    add_label_detail_table(doc, [
        ("موقع الزبون", "apiFetch يطبق مهلة افتراضية 7 ثوانٍ، يستخدم AbortController، يحول الجسم غير FormData إلى JSON، ويرجع data مباشرة. مسارا الاستعادة لا يرسلان التوكن القديم."),
        ("لوحة الموظفين", "TAZA.Http يوفر get/post/put/patch/delete/upload؛ عند 401 في طلب محمي يعرض انتهاء الجلسة ثم يمسح Local Storage ويعيد إلى الدخول، بينما 401 تسجيل الدخول يعرض خطأ الاعتماد فقط."),
        ("رفع الملفات", "لا يضبط Content-Type يدويًا؛ يترك المتصفح يضيف multipart boundary ويرسل Bearer Token."),
        ("Base URL", "إذا كانت الواجهة على منفذ تطوير محلي تستخدم Laravel على المنفذ 8000، وإلا تستخدم origin نفسه؛ يمكن override في بيئة التطوير فقط."),
    ])

    add_heading(doc, "2.3 المصادقة والأدوار", 2)
    add_body(doc, "تسجيل دخول الموظف يحذف توكناته القديمة وينشئ توكنًا بقدرات الدور وصلاحية افتراضية 480 دقيقة. تسجيل دخول الزبون يستخدم Sanctum كذلك. TrackCustomerIp يسجل عنوان IP في كل مسار زبون محمي لدعم الحماية والحظر.")
    add_label_detail_table(doc, [
        ("general_manager", "الموظفون، الزبائن، كل الطلبات والتقارير، إعدادات المطعم، الولاء والإشعارات."),
        ("order_manager", "الطلبات العادية والحجوزات وتغيير الحالات والأرشفة وإشعار الزبون."),
        ("delivery_manager", "التوصيلات، السائقون، الإسناد، الإلغاء، الإعدادات."),
        ("driver", "التوصيلات المسندة، بدء التوصيل، التسليم، وإشعار الزبون."),
        ("inventory_manager", "المنتجات، المخزون، الأسعار، الصور، العروض وربط منتجاتها."),
        ("finance_manager", "حسابات التحصيل، حركات الدفع، الاسترداد والتقارير المالية."),
        ("communication_manager", "معلومات المطعم والمعرض واقتراحات الوجبات والمراجعات وتقارير AI."),
    ], header=("الدور", "نطاق العمليات"))

    add_heading(doc, "3. عمليات موقع الزبون طرفًا إلى طرف", 1)
    add_heading(doc, "3.1 التدفقات التفاعلية الرئيسية", 2)
    customer_flows = [
        {"name":"تسجيل حساب جديد", "trigger":"إرسال نموذج التسجيل في register.html.", "client":"router-auth.js يجمع الاسم والبريد والهاتف وكلمة المرور والتأكيد، ويتحقق من الحقول والصورة الاختيارية محليًا.", "request":"يرسل POST /api/customer/auth/register؛ عند وجود صورة قد يعقب إنشاء الحساب رفع الصورة وفق تدفق الصفحة.", "server":"CustomerAuthController@register يطبّع المدخلات، يطبق قواعد الأمان والتفرد، ينشئ Customer وحساب LoyaltyAccount داخل معاملة ثم ينشئ توكن Sanctum.", "result":"setAuth يحفظ التوكن والمستخدم، تظهر رسالة نجاح، ثم ينتقل المستخدم إلى صفحته المطلوبة/الرئيسية.", "errors":"422 للحقول أو التكرار، و500 مع رسالة آمنة عند فشل المعاملة."},
        {"name":"تسجيل دخول الزبون", "trigger":"زر تسجيل الدخول في login.html.", "client":"تعطّل الواجهة زر الإرسال وتقرأ البريد وكلمة المرور.", "request":"POST /api/customer/auth/login بجسم JSON.", "server":"يتحقق CustomerAuthController من الاعتماد والحظر، يسجل IP، وينشئ توكنًا جديدًا ويرجع customer.", "result":"يحفظ setAuth البيانات ويعيد المستخدم إلى الصفحة المقصودة؛ عند الفشل يعاد تفعيل الزر وتعرض الرسالة.", "errors":"401 لاعتماد خاطئ، 403 للحساب الموقوف، و429 عند تجاوز حد المحاولات."},
        {"name":"نسيت كلمة المرور وإعادة الضبط", "trigger":"إرسال البريد ثم فتح رابط الاستعادة وإرسال كلمة المرور الجديدة.", "client":"router-auth.js يستدعي forgot-password، ثم reset-password باستخدام email وtoken من الرابط وكلمة المرور المؤكدة.", "request":"POST /api/customer/auth/forgot-password ثم POST /api/customer/auth/reset-password؛ لا يرسل التوكن القديم معهما.", "server":"يستخدم Password broker؛ الطلب الأول يعيد رسالة عامة لمنع كشف وجود الحساب، والثاني يتحقق من الرمز ويحدّث hash ويلغي كل توكنات الحساب القديمة.", "result":"رسالة نجاح ثم العودة إلى تسجيل الدخول.", "errors":"422 لرابط غير صالح/منتهي، 429 للتقييد، 500 لفشل حفظ غير متوقع."},
        {"name":"تحميل القائمة والعروض والبيانات الحية", "trigger":"فتح الرئيسية أو menu.html وتغيير البحث/التصنيف.", "client":"تحمّل state/catalog البيانات من /public/live-data أو نقاط المنتجات والعروض وتوحّدها، ثم تطبق الفلاتر محليًا وتعيد رسم البطاقات.", "request":"GET /api/public/live-data أو GET /api/public/products و/public/offers مع معاملات البحث والتصنيف.", "server":"RestaurantController وProductController وOfferController يعيدون العناصر المتاحة فقط والأسعار والمخزون والصور وإعدادات المطعم.", "result":"تظهر البطاقات والحالة المفتوح/المغلق؛ عند تعذر API تستخدم الواجهة بيانات fallback للعرض وتمنع العمليات التي تتطلب بيانات حية.", "errors":"مهلة/انقطاع يعرض تنبيهًا آمنًا ولا يسمح بادعاء نجاح طلب حقيقي."},
        {"name":"إدارة السلة", "trigger":"إضافة عنصر، زيادة/تقليل الكمية، حذف عنصر، أو اختيار نوع الطلب.", "client":"customer-cart.js يتحقق من التوفر والحد الأقصى، يخزن العناصر والملاحظات في AppState/Local Storage ويطلق taza:cart-updated؛ لا يوجد طلب API في هذه المرحلة.", "request":"لا يوجد طلب حتى بدء checkout؛ التغيير محلي وقابل للمراجعة.", "server":"لا ينفذ الخادم شيئًا؛ عند المزامنة التالية يصالح reconcileCartWithCatalog العناصر مع الكتالوج الحي.", "result":"يتحدث العدد والإجمالي فورًا؛ العناصر المنتهية تبقى بعلامة غير متوفر ويجب حذفها للمتابعة.", "errors":"يمنع تجاوز المخزون أو المتابعة والمطعم مغلق أو النوع غير مختار."},
        {"name":"تحديد موقع التوصيل وحساب السعر", "trigger":"اختيار عنوان محفوظ أو النقر على الخريطة/تحديد الموقع في delivery.html.", "client":"delivery.js يهيئ الخريطة، يثبت الإحداثيات، ويطلب quote؛ يرسم GeoJSON القادم ويعرض المسافة والمدة والتكلفة.", "request":"GET /api/public/delivery/quote?latitude=...&longitude=....", "server":"RestaurantController@deliveryQuote يستدعي DeliveryRouteService؛ الخدمة تطلب OSRM مع timeout وcache، وتعود إلى Haversine وخط مستقيم عند التعذر. ثم تفحص الحد وتحسب التكلفة.", "result":"يحفظ الاختيار في AppState ويعرض provider وis_fallback ضمنيًا في حالة المسار، ويتيح المتابعة إذا كان ضمن النطاق.", "errors":"422 للإحداثيات، خارج النطاق يعيد is_within_range=false وتكلفة null، وفشل OSRM لا يفشل العملية بل يستخدم البديل المعلّم."},
        {"name":"التحقق من الطاولة", "trigger":"اختيار وقت ورقم طاولة في reservation.html.", "client":"reservation.js يطلب كتالوج الطاولات ثم يعيد الفحص عند تغير الوقت أو الطاولة ويحدث الخريطة المرئية والتكلفة.", "request":"GET /api/public/reservations/tables?reservation_time=... ثم GET /api/public/reservations/table/{n}/availability.", "server":"ReservationController يكمل تلقائيًا الحجوزات المنتهية، يمنع الماضي وما بعد 24 ساعة، يبحث عن تعارض لمدة 60 دقيقة، ويحسب VIP والمقاعد الزائدة.", "result":"تظهر available/reserved والتكلفة الإضافية؛ لا يمكن المتابعة عند التعارض.", "errors":"404 لطاولة غير موجودة، 422 لوقت غير صالح، أو رد متاح=false عند التعارض."},
        {"name":"إنشاء الطلب بأنواعه الثلاثة", "trigger":"زر تأكيد الدفع/الطلب بعد اكتمال السلة وبيانات النوع.", "client":"payment.js يبني buildOrderPayload: type وnotes وitems، ويضيف العنوان والإحداثيات للتوصيل أو الطاولة/الوقت/المقاعد للحجز.", "request":"POST /api/customer/orders مع Bearer Token.", "server":"OrderController@customerStore يتحقق من الزبون والحظر وفتح المطعم والمدخلات. يعيد حساب الأسعار من قاعدة البيانات، يفحص المخزون والعرض، يحسب route أو الحجز، ثم ينشئ Order وOrderItem وDeliveryOrder/ReservationOrder ويخصم المخزون ويرسل إشعار مدير الطلبات داخل معاملة.", "result":"يرجع order وملخص items_total/extra_cost/final_price بكود 201، ثم تبدأ الواجهة خطوة الدفع.", "errors":"423 لإغلاق المطعم، 422 للبيانات، رفض التوفر/المخزون/نطاق التوصيل/تعارض الطاولة، وrollback كامل عند أي استثناء."},
        {"name":"دفع الطلب", "trigger":"اختيار الطريقة والضغط على تأكيد الدفع.", "client":"بعد نجاح إنشاء الطلب يستدعي payment.js الدفع باستخدام معرف الطلب والطريقة والقيمة اللازمة.", "request":"POST /api/customer/orders/{id}/pay بجسم method وربما phone/pin_code أو points_required.", "server":"PaymentController يرفض الدفع المكرر والحالات النهائية، يعطل Syriatel/Sham حاليًا بكود 423، ويستدعي PaymentRecord::processPayment داخل معاملة. نقاط الولاء تُخصم عند الاستبدال، وتُمنح فورًا للدفع الإلكتروني/الاختباري، بينما الكاش يبقى pending وتُمنح نقاطه عند اكتمال الخدمة.", "result":"يمسح العميل السلة ويعرض رقم الطلب والطريقة والنقاط المكتسبة وحالة الطلب.", "errors":"رصيد نقاط غير كاف، سجل دفع قائم، طريقة معطلة، أو rollback ورسالة 500 عند الاستثناء."},
        {"name":"متابعة الطلبات والإلغاء", "trigger":"فتح orders.html أو الضغط على إلغاء طلب معلق.", "client":"الحساب يحمّل GET /customer/orders، يعرض customer_status الموحد والخطوات، ويطلب تأكيد الإلغاء قبل الإرسال.", "request":"GET /api/customer/orders وDELETE /api/customer/orders/{id} مع سبب اختياري.", "server":"OrderController يقيد البيانات بمالكها. الإلغاء مسموح للزبون في pending فقط ويستدعي OrderCancellationService بقفل قاعدة البيانات لمنع التكرار؛ يعيد المخزون والمال/النقاط مرة واحدة، يحدث الحالات ويرسل إشعارات وقد يطبق الحظر التلقائي للإلغاءات المتكررة.", "result":"تتحدث البطاقة ويعرض refund وalready_cancelled؛ يعاد تحميل البيانات.", "errors":"403 لغير المالك، 404، أو رفض انتقال الحالة بعد بدء المعالجة."},
        {"name":"تقييم السائق والوجبات", "trigger":"زر التقييم الذي يظهر فقط بعد التسليم/الاكتمال.", "client":"يفتح account.js نافذة النجوم والملاحظة ثم يرسل التقييم ويغلق الزر بعد النجاح.", "request":"POST /api/customer/delivery/{id}/rate أو POST /api/customer/orders/{orderId}/products/{productId}/rate.", "server":"ReviewController يتحقق من ملكية الطلب والحالة النهائية ومنع التكرار، ثم ينشئ Review أو يحدث حقول تقييم التوصيل.", "result":"يعرض التقييم في الطلب وتنعكس البيانات في لوحات المدير المختص.", "errors":"رفض قبل اكتمال الخدمة، منتج ليس ضمن الطلب، تقييم خارج 1–5، أو تقييم مكرر."},
        {"name":"تعديل الملف والعناوين والصورة", "trigger":"حفظ بيانات الحساب، مزامنة عناوين المنزل/العمل، أو رفع الصورة.", "client":"account.js يجمع التغييرات ويطلب كلمة المرور للعمليات الحساسة، ويستخدم FormData للصورة.", "request":"PUT /api/customer/profile، GET/PUT/DELETE /api/customer/saved-addresses، وPOST /api/customer/avatar.", "server":"المتحكمات تطبع النصوص وتتحقق من التفرد وكلمة المرور وحدود العنوان والإحداثيات والصورة، ثم تحفظ وتعيد التفاصيل.", "result":"تحدث AppState والمظهر والعناوين المستخدمة في التوصيل.", "errors":"422 لكلمة مرور أو ملف/عنوان غير صالح، 403/401 للجلسة، و404 لنوع عنوان غير موجود."},
        {"name":"المساعد الذكي واقتراح وجبة", "trigger":"إرسال رسالة للمساعد أو إرسال اقتراح وجبة مع صورة اختيارية.", "client":"ai.js يحافظ على conversation_id، يعرض حالة الكتابة والردود السريعة والمنتجات؛ ويرفع الاقتراح عبر FormData ثم يراقب حالته.", "request":"POST /api/customer/ai/chat، GET /customer/ai/history، POST /customer/meal-suggestion، GET /customer/meal-suggestions.", "server":"AIController يضمن أن سياق المحادثة يخص الزبون ويستدعي AiConversation/GenerativeMealAdvisor؛ MealSuggestionController يمنع التكرار خلال 24 ساعة ويخزن الصورة ويرسل إشعارًا لمدير التواصل.", "result":"يرجع reply وintent وsuggested_items وquick_replies، أو سجل اقتراح بحالة pending ثم حالات المراجعة اللاحقة.", "errors":"422 لنص قصير/طويل أو صورة غير صالحة، 403 للحساب الموقوف، واستجابة بديلة حتمية إذا لم تتوفر خدمة OpenAI."},
        {"name":"الإشعارات", "trigger":"فتح صفحة الإشعارات أو الضغط على إشعار/تعليم الكل مقروءًا.", "client":"تحمل القائمة وتحدث عداد unread وتزيل النمط unread محليًا بعد نجاح الطلب.", "request":"GET /api/customer/notifications، PUT /read-all، PUT /{id}/read.", "server":"NotificationController يقيد النتائج بالمستلم، ويعيد مجموعات وإحصاءات ويحدث is_read/read_at.", "result":"تتحدث الشارة والقائمة ويمكن الانتقال إلى الطلب المرتبط من data.", "errors":"لا يسمح بتعديل إشعار لا يخص المستخدم؛ تعرض الواجهة رسالة آمنة وتعيد المزامنة."},
    ]
    for flow in customer_flows:
        add_flow(doc, flow)

    add_heading(doc, "4. العمليات المشتركة في لوحة الموظفين", 1)
    shared_dashboard = [
        ("تسجيل الدخول", "index.html يرسل POST /api/auth/employee/login؛ يخزن Auth.save التوكن وemployee، ثم Auth.redirectToHome يختار الصفحة حسب role."),
        ("حماية الصفحة", "initDashboardPage يستدعي Auth.requireRole؛ غياب الجلسة يعيد للدخول، والدور الخاطئ يعاد توجيهه لصفحته الصحيحة."),
        ("تحديث الملف", "employee-profile.js يحمّل /auth/employee/me ويحفظ PUT /profile بعد طلب current_password؛ الصورة POST/DELETE /avatar."),
        ("الإشعارات", "NotifBadge يستعلم unread-count دوريًا؛ اللوحة وصفحات الأدوار تستخدم list/read/read-all وتحدث الشارة."),
        ("التقارير المشتركة", "كل موظف يقرأ ويرسل ويراجع ويؤرشف تقاريره عبر /employee/reports؛ صلاحية العرض تقيد بالمُرسل/المستلم."),
        ("المظهر واللغة", "عمليات محلية: تحفظ اللغة والثيم، تحدث اتجاه الصفحة، وتعيد تلوين المخططات دون API."),
        ("التحديث الحي", "LiveSync والتحديثات الدورية تعيد طلب بيانات التبويب المفتوح فقط وتتوقف/تخف عند عدم ظهور الصفحة."),
        ("تسجيل الخروج", "نافذة تأكيد ثم POST /auth/employee/logout؛ حتى لو فشل الطلب محليًا تمسح Auth.logout التخزين وتعود للدخول."),
    ]
    add_label_detail_table(doc, shared_dashboard, header=("العملية", "التنفيذ من الواجهة إلى النتيجة"))

    role_sections = [
        ("5. عمليات المدير العام", [
            ("نظرة عامة", "طلبات متوازية لـ /auth/employee/me و/admin/orders/stats و/admin/reports/stats وغيرها؛ تُحوّل النتائج إلى بطاقات ومخططات مع تحمل الفشل الجزئي."),
            ("إدارة الموظفين", "عرض وبحث/فلترة، إنشاء POST، تعديل PUT/PATCH مع POST fallback، تفعيل/تعطيل أو حذف DELETE، رفع/حذف الصورة، إرسال إشعار، وعرض/إضافة التقييمات."),
            ("إدارة الزبائن", "فلاتر most_orders/top_spenders/top_loyalty/suspicious، عرض الطلبات، تحذير، حظر/رفع الحظر، وإشعار جماعي queued."),
            ("مراقبة الطلبات", "GET /admin/orders مع type/status/date، إحصاءات وتفاصيل كاملة دون تغيير تشغيلي مباشر."),
            ("التقارير", "قائمة الوارد وإحصاءاته، فتح التفاصيل، وإرسال تعليمات عبر POST /admin/reports/{id}/send."),
            ("الولاء", "عرض الحسابات والإحصاءات والحركات، تعديل النقاط مع سبب، وتحديث أسعار/مضاعفات النظام ضمن الإعدادات."),
            ("إعدادات المطعم", "تحديث تسعير التوصيل وحده وإحداثيات الأصل، إعدادات الحجز، فتح/إغلاق المطعم، ورفع الشعار."),
        ]),
        ("6. عمليات مدير الطلبات والحجوزات", [
            ("صندوق الطلبات", "GET /orders و/pending و/normal؛ الفلاتر محلية أو query، وتُبنى بطاقات مع العمر والحالة والدفع والملاحظات."),
            ("تغيير حالة الطلب", "زر الإجراء يحدد الحالة التالية ثم PUT /orders/{id}/status. Order::STATUS_FLOW يمنع القفز؛ الانتقال إلى completed في طلب توصيل يرسل إشعارًا لمدير التوصيل."),
            ("إشعار الزبون", "نافذة نص ثم POST /orders/{id}/notify-customer؛ ينشأ Notification مرتبط بالطلب."),
            ("إدارة السجل", "بعد إغلاق العملية فقط: archive، restore، delete؛ الخادم يتحقق من isOperationallyClosed قبل إدارة السجل."),
            ("الحجوزات", "تبويبات اليوم/الكل/القادمة وخريطة الطاولات؛ PUT status إلى confirmed/seated/completed/cancelled/no_show. التأكيد يعيد فحص التوفر، والجلوس ينتظر اكتمال تجهيز الطلب."),
            ("إلغاء من الموظف", "يمر عبر OrderCancellationService نفسه، يعكس الموارد مرة واحدة ويرجع refund وalready_cancelled."),
        ]),
        ("7. عمليات مدير التوصيل والسائق", [
            ("لوحة التوصيل", "GET /delivery/stats و/active و/drivers؛ ترتيب الأولويات يظهر الطلب الجاهز بلا سائق أولًا."),
            ("إسناد سائق", "يفتح modal بالسائقين المتاحين ثم PUT /delivery/{id}/assign. لا يقبل إلا طلبًا أساسيًا completed وتوصيلًا pending وسائقًا نشطًا؛ يتحول مباشرة إلى in_delivery وتصل إشعارات للسائق والزبون."),
            ("إلغاء التوصيل", "مدير التوصيل فقط يرسل status=cancelled وسببًا؛ الخادم يسوي الطلب والدفع والمخزون والنقاط."),
            ("إعدادات التوصيل", "GET/PUT /delivery/settings؛ تحديث التكلفة لكل كم والحد الأقصى وإحداثيات المطعم، مع حاسبة محلية للمعاينة."),
            ("عمل السائق", "GET /delivery/active أو /assigned مقيد بالسائق؛ PUT status يسمح in_delivery ثم delivered فقط للسائق المسند. delivered يسجل الوقت ويكمل Order ويمنح نقاط الكاش."),
            ("إشعار السائق للزبون", "POST /delivery/{id}/notify-customer متاح للسائق المسند فقط ويقيد الرسالة بـ500 حرف."),
            ("السجل والتقييمات", "GET /delivery مع فلاتر route_quality/status، وGET /driver/{id}/ratings|stats لعرض الأداء."),
        ]),
        ("8. عمليات مدير المخزون والعروض", [
            ("المنتجات", "عرض وفلاتر all/low/out/stats؛ إضافة وتعديل، سعر، سعر ولاء، مخزون، تفعيل، حذف، ورفع صورة."),
            ("حفظ المنتج", "saveProduct يتحقق من الحقول والصورة؛ ينشئ POST أو يحدث PUT، ثم يرفع الصورة منفصلة إذا وجدت ويعيد تحميل الشبكة."),
            ("تحديث المخزون", "PATCH /products/{id}/stock؛ Product يطبق الكمية ويرسل تنبيه نفاد/انخفاض مع deduplication عند اللزوم."),
            ("بلاغ عدم التوفر", "الزبون يرسل /customer/products/{id}/report-unavailable؛ الخادم يرسل إشعارًا لمدير المخزون وتمنع الواجهة التكرار في الجلسة."),
            ("العروض", "إنشاء/تعديل وحذف وتفعيل، إضافة/إزالة منتجات وكمياتها، سعر ولاء وصورة، وفلاتر active/expired/upcoming."),
            ("بث العرض", "POST /offers/{id}/broadcast يضع BroadcastNewOffer في الطابور بدل إنشاء كل الإشعارات داخل طلب المتصفح."),
        ]),
        ("9. عمليات المدير المالي", [
            ("حسابات التحصيل", "عرض الحسابات والملخص، إنشاء/تعديل، تعيين أساسي، تعديل رصيد، سحب، وحذف. كل عملية تفحص النوع والنشاط والسعة/الرصيد."),
            ("حركات الدفع", "GET /finance/payments مع الفلاتر والإحصاءات؛ فتح سجل يعيد الطلب والزبون والحساب."),
            ("الاسترداد", "POST /finance/payments/{id}/refund؛ يمنع التكرار ويحدث status=refunded ويعكس الرصيد الداخلي حسب الطريقة."),
            ("التقرير المالي", "زر التوليد يطلب GET /finance/report؛ المتحكم يجمع الفترة ويضع GenerateFinancialReport في طابور reports، ثم يصل Report للمدير العام."),
            ("طرق الدفع الفعلية", "cash وloyalty_points يعملان؛ test_payment للتقييم فقط إذا سمحت البيئة؛ Syriatel Cash وSham Cash مرفوضتان حاليًا 423 ولا يوجد مزود خارجي متصل."),
        ]),
        ("10. عمليات مدير التواصل والذكاء الاصطناعي", [
            ("معلومات المطعم", "GET/PUT /communication/restaurant لتعديل الاسم والوصف والتواصل والساعات والروابط ومحتوى About."),
            ("المعرض", "رفع عدة صور بالتتابع مع تقدم، تعديل caption/type، تغيير order، toggle الظهور، وحذف السجل والملف."),
            ("اقتراحات الوجبات", "عرض وإحصاءات وتفاصيل؛ review يضيف ملاحظات، implement أو reject يغير الحالة ويرسل إشعارًا للزبون."),
            ("تقييمات الزبائن", "GET /reviews/customers يعرض تقييمات المنتجات للمتابعة والتحليل."),
            ("محادثات AI", "GET /ai/conversations و/stats للفلاتر والمؤشرات؛ POST /ai/generate-report يضع GenerateDailyAiReport في الطابور."),
            ("تقارير AI", "GET /communication/ai-reports ثم POST /{id}/forward لتمرير التقرير للمدير العام مع رسالة."),
        ]),
    ]
    for title, rows in role_sections:
        add_heading(doc, title, 1)
        for name, detail in rows:
            add_heading(doc, name, 2)
            add_body(doc, detail)

    add_heading(doc, "11. المهام الخلفية والخدمات الخارجية والنسخ الاحتياطي", 1)
    add_label_detail_table(doc, [
        ("BroadcastNewProduct", "يرسل إعلان منتج جديد للزبائن عبر الطابور مع إنشاء إشعارات مجمعة."),
        ("BroadcastNewOffer", "يرسل بث عرض للزبائن ويمنع إبطاء استجابة زر البث."),
        ("BroadcastCustomerAnnouncement", "يعالج الإشعار الجماعي من المدير العام."),
        ("GenerateFinancialReport", "ينشئ تقريرًا ماليًا للفترة ويرسله كسجل Report للمدير العام."),
        ("GenerateDailyAiReport", "يجمع محادثات/اقتراحات AI وينشئ تقريرًا لمدير التواصل."),
        ("OSRM", "DeliveryRouteService يطلب /route/v1/{profile}/{coordinates} مع cache 10 دقائق افتراضيًا وtimeouts، ثم fallback Haversine عند الفشل."),
        ("OpenAI Responses", "GenerativeMealAdvisor يستخدم نموذجًا مضبوطًا في services.openai، مخطط JSON منضبطًا، ويقبل فقط معرفات منتجات موجودة؛ يبقى الرد الحتمي المحلي fallback."),
        ("البريد", "استعادة كلمة المرور تستخدم Password broker؛ الفشل لا يكشف وجود الحساب ويحافظ على الرمز السابق إن وجد."),
        ("النسخ الاحتياطي", "taza:backup ينشئ snapshot وقائمة وchecksums؛ taza:backup-verify يختبر الاستعادة معزولة؛ taza:backup-restore ينشئ نسخة أمان قبل الاستعادة ويدخل الصيانة ويتراجع عند الفشل."),
    ], header=("المكون", "كيف يعمل"))
    add_code(doc, "php artisan queue:work --queue=mail,notifications,reports,default\nphp artisan queue:monitor mail,notifications,reports,default --max=100\nphp artisan queue:failed\nphp artisan taza:backup\nphp artisan taza:backup-verify <backup-id>")

    add_heading(doc, "12. خرائط الحالات والآثار المتقاطعة", 1)
    add_heading(doc, "12.1 الطلب الأساسي", 2)
    add_body(doc, "pending ← confirmed ← ready ← completed، مع cancelled وفق الدور والحالة. يستخدم Order::canChangeStatus وchangeStatus لمنع الانتقالات غير المعرفة. عند اكتمال الطلب النقدي يصبح سجل الدفع completed وتمنح نقاط الولاء مرة واحدة.")
    add_heading(doc, "12.2 التوصيل", 2)
    add_body(doc, "pending ← in_delivery ← delivered. الحالات القديمة assigned وpicked_up تطبع إلى in_delivery في تجربة العرض. الإسناد يرسل التوصيل مباشرة إلى in_delivery؛ delivered يكمل الطلب الأساسي ويسجل actual_delivery_time.")
    add_heading(doc, "12.3 الحجز", 2)
    add_body(doc, "pending ← confirmed ← seated ← completed، مع cancelled أو no_show. مدة الجلسة 60 دقيقة؛ الفحص اللاحق يحول الحجز المنتهي إلى completed إذا كان seated أو no_show خلاف ذلك ويحرر الطاولة منطقيًا.")
    add_heading(doc, "12.4 الإلغاء الموحد", 2)
    cancellation_num = new_numbering(doc)
    for item in (
        "يقفل Order بقفل for update داخل DB transaction لمنع سباق طلبين.",
        "إذا كان الطلب ملغى يعيد ملخصًا مع already_cancelled=true ولا يكرر أي عكس.",
        "يعيد كميات المنتجات المباشرة ومكونات العروض إلى المخزون.",
        "يسوي الدفع: استرداد أو إلغاء pending حسب الطريقة، ويعيد نقاط الاستبدال عند الحاجة.",
        "يعكس نقاط الكسب المتعلقة بالطلب إن كانت قد منحت.",
        "يحدّث Order وDeliveryOrder/ReservationOrder ويرسل إشعارات الإلغاء.",
    ):
        add_numbered(doc, item, cancellation_num)
    add_callout(doc, "ملاحظة مالية", "إلغاء زر الطلب وإلغاء الحجز وإلغاء التوصيل جميعها تتقاطع مع OrderCancellationService، لذلك لا ينبغي لأي واجهة مستقبلية تنفيذ عكس المخزون أو النقاط محليًا أو بمسار جانبي.", fill="FFF7E6", accent=AMBER)

    doc.add_page_break()
    add_heading(doc, f"13. ملحق API الكامل ({len(routes)} مسارًا)", 1)
    add_body(doc, "هذا الملحق يغطي كل مسار ظاهر في قائمة Laravel وقت إنشاء المستند. الجدول يذكر طريقة HTTP والمسار والمتحكم والغرض. التفاصيل الدقيقة للـpayload وقواعد التحقق تبقى مرجعيتها النهائية في المتحكم المذكور.")
    add_route_catalog(doc, routes)

    add_heading(doc, "14. دليل الملفات المرجعية", 1)
    add_label_detail_table(doc, [
        ("routes/api.php", "تعريف كل مسارات API وترتيبها والمجموعات المحمية."),
        ("public/frontend/assets/js/core/state.js", "AppState وapiFetch والمهلة والغلاف العام للأخطاء."),
        ("public/frontend/assets/js/core/customer-cart.js", "السلة والمصالحة مع الكتالوج والانتقال إلى checkout."),
        ("public/frontend/assets/js/pages/*.js", "تدفقات التسجيل والحساب والقائمة والتوصيل والحجز والدفع والذكاء الاصطناعي."),
        ("public/dashboard/assets/js/config.js", "API map وAuth وHttp والإشعارات والتهيئة المشتركة."),
        ("public/dashboard/assets/js/pages/**", "عمليات كل دور وأحداث الأزرار وعرض النتائج."),
        ("app/Http/Controllers/API", "التحقق من الدور والمدخلات وتنسيق الردود وتنسيق الخدمات."),
        ("app/Models", "قواعد المجال والحالات والعلاقات والآثار."),
        ("app/Services/OrderCancellationService.php", "الإلغاء idempotent وتسوية الموارد."),
        ("app/Services/DeliveryRouteService.php", "OSRM والـcache والـfallback."),
        ("app/Services/GenerativeMealAdvisor.php", "الدمج مع OpenAI وضبط المخرجات والبديل."),
        ("app/Jobs", "الإشعارات الجماعية والتقارير المؤجلة."),
        ("tests/Feature", "اختبارات المصادقة والطلبات والدفع والحجز والتوصيل والأمان والطوابير والنسخ."),
    ], header=("المسار", "ما الذي يثبته"))
    add_callout(doc, "خلاصة الصيانة", "عند تعديل زر أو endpoint: حدّث route وAPI map ودالة الصفحة والمتحكم وقواعد الحالة والاختبار المرتبط معًا. أكثر مناطق الخطر هي إنشاء الطلب، الإلغاء، الدفع، المخزون، الولاء، وحالات التوصيل/الحجز.", fill="EEF6F2", accent=GREEN)

    # Core metadata and compatibility settings.
    doc.core_properties.title = "TAZA 041 — الدليل التقني الشامل للعمليات وواجهات API"
    doc.core_properties.subject = "توثيق طرف إلى طرف لموقع الزبون ولوحة الموظفين وLaravel API"
    doc.core_properties.author = "TAZA 041 Technical Documentation"
    doc.core_properties.keywords = "TAZA041, API, Laravel, operations, Arabic, dashboard"
    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(str(OUTPUT))
    print(f"routes={len(routes)} paragraphs={len(doc.paragraphs)} tables={len(doc.tables)}")


if __name__ == "__main__":
    main()
